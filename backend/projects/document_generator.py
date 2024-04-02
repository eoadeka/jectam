from docx import Document
from io import BytesIO
import json
from .models import Project,  Document as Doc
from django.core.files.base import ContentFile


def generate_document(json_data, project, file_type):
    project = Project.objects.get(project_id=project)

    # Retrieve the latest version number for the project
    try:
        latest_version = Doc.objects.filter(project=project, document_type=file_type).latest('created_at')
        version_number = latest_version.version_number + 1
    except Doc.DoesNotExist:
        version_number = 1  # If no previous version exists, start from 1
    
     # Get the corresponding value from TEMPLATE_TYPES for the provided key
    file_type_value = next((value for key, value in Doc.TEMPLATE_TYPES if key == file_type), None)

    # Update the latest version number in the database
    doc_title = f"{project.title}_{file_type}_v{version_number}"
    new_doc = Doc.objects.create(project=project, title=doc_title, version_number=version_number, document_type=file_type)

    print(new_doc.document_id)

    document = Document()

    # print(document)

    # Add document title
    document.add_heading(json_data["title"], level=0)
    records = (
        ('Project Name', project.title),
        ('Date', new_doc.created_at.strftime('%Y-%m-%d %H:%M:%S')),
        ('Author', "1"),
        ('Document Number', f"{project.title}_v{version_number}"),
        ('File Type', file_type_value) 
    )
    table = document.add_table(rows=len(records), cols=2)  # Set rows to the length of records
    table.style = 'Table Grid'
    for index, (key, value) in enumerate(records):  # Use enumerate to access index
        row_cells = table.rows[index].cells
        row_cells[0].text = key
        row_cells[1].text = value

    document.add_page_break()

   # Add sections
    sections = []
    for section in json_data["sections"]:
        section_data = {
            "title": section["title"],
            "content": section["content"]
        }
        sections.append(section_data)

        document.add_heading(section["title"], level=2)
        document.add_paragraph(section["content"])

    # Serialize document content as JSON
    document_content = {
        "records": records,
        "sections": sections
    }

    # Update the new Doc instance with the generated content
    new_doc.content = json.dumps(document_content)
    new_doc.save()
    # print(document)

    return document,  new_doc.document_id

